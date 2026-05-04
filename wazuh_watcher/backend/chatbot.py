"""
chatbot.py — Gemini-powered security analyst chatbot for Wazuh-ATHEA.

The chatbot receives the live pipeline snapshot (anomalies, normal alerts,
SHAP explanations, MITRE stages, cluster IDs) and answers analyst questions
by correlating anomalous vs normal behaviour.
"""

import logging
import os
import io
import re
from typing import Any

import google.generativeai as genai
from dotenv import load_dotenv
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

load_dotenv()

logger = logging.getLogger(__name__)

# ── Gemini configuration ────────────────────────────────────────────────────
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GEMINI_MODEL   = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-preview-04-17")

if GEMINI_API_KEY and GEMINI_API_KEY != "your_gemini_api_key_here":
    genai.configure(api_key=GEMINI_API_KEY)
    _gemini_ready = True
else:
    logger.warning("GEMINI_API_KEY not configured — chatbot will return placeholder responses.")
    _gemini_ready = False


# ── Request / Response models ───────────────────────────────────────────────
class ChatMessage(BaseModel):
    role: str    # "user" or "model"
    content: str


class ChatRequest(BaseModel):
    message: str
    history: list[ChatMessage] = []


class ChatResponse(BaseModel):
    reply: str
    model: str
    context_loaded: bool


class ReportResponse(BaseModel):
    report: str


class DocxRequest(BaseModel):
    report_text: str


# ── Feature name → human-readable description ───────────────────────────────
_FEATURE_DESCRIPTIONS: dict[str, str] = {
    "IsFailedLogin":        "failed authentication attempt",
    "IsProcessCreation":    "process creation event",
    "IsNetworkConnection":  "network connection event",
    "IsPrivilegeUse":       "privilege use / escalation",
    "IsScheduledTask":      "scheduled task creation/modification",
    "IsLogCleared":         "security log cleared (defense evasion)",
    "IsServiceInstall":     "service installation (persistence)",
    "IsLateralMove":        "lateral movement indicator",
    "IsRegistryMod":        "registry modification",
    "IsPowerShell":         "PowerShell activity",
    "LogHasSuspicious":     "suspicious keywords (mimikatz/bypass/shellcode/etc.)",
    "LogHasBase64":         "base64-encoded content in log",
    "LogHasPowerShell":     "PowerShell strings in log text",
    "IsNightTime":          "occurred during off-hours (before 6AM or after 10PM)",
    "IsWeekend":            "occurred on a weekend",
    "FailedLoginRatio":     "ratio of failed logins in the rolling time window",
    "HighSeverityRatio":    "ratio of high-severity (level ≥10) events in the window",
    "EventsPerMinute":      "event frequency in the rolling window",
    "UniqueEventIDsInWindow":"variety of different Event IDs in the rolling window",
    "TimeSincePrevEvent":   "time gap since previous event from same agent",
    "EventIDThreatTier":    "threat tier classification of this Event ID (0=benign, 3=high-risk)",
    "EventIDRarity":        "how rare this Event ID is in the current batch",
    "MitreTacticWeight":    "MITRE ATT&CK tactic severity weight",
    "MitreTacticCount":     "number of MITRE tactics mapped to this event",
    "RuleLevel":            "Wazuh rule severity level (0-15)",
    "RuleLevelNorm":        "normalised Wazuh rule level",
    "CommandLineLength":    "length of the command line (longer = more suspicious)",
    "LogLength":            "raw log length",
    "LogEntropy":           "Shannon entropy of the log (high = obfuscated/encoded content)",
    "CommandLineEntropy":   "Shannon entropy of the command line",
    "ComputerEventCount":   "total number of events from this agent in the batch",
    "Hour":                 "hour of day (0-23)",
    "DayOfWeek":            "day of week (0=Monday, 6=Sunday)",
    "HourSin":              "cyclic hour encoding (sine)",
    "HourCos":              "cyclic hour encoding (cosine)",
    "MitrePresent":         "whether a MITRE ATT&CK ID was mapped to this event",
}


def _describe_feature(feature: str) -> str:
    """Return a human-readable description for a SHAP feature name."""
    return _FEATURE_DESCRIPTIONS.get(feature, feature.replace("_", " ").lower())


def _format_shap_contributions(shap_values: list[dict]) -> str:
    """Convert raw SHAP dicts into human-readable text for the LLM."""
    if not shap_values:
        return "  (no SHAP data available)"

    lines = []
    for sv in shap_values:
        feature   = sv.get("feature", "unknown")
        shap_val  = sv.get("shap_value", 0.0)
        alert_val = sv.get("alert_value", 0.0)
        direction = sv.get("direction", "unknown")

        direction_str = "↑ pushed toward ANOMALY" if shap_val > 0 else "↓ pushed toward NORMAL"
        desc = _describe_feature(feature)
        lines.append(
            f"  • [{feature}] = {alert_val} → {direction_str} "
            f"(SHAP impact: {shap_val:+.4f}) — {desc}"
        )
    return "\n".join(lines)


def _format_alert_for_prompt(alert: dict, is_anomaly: bool, index: int) -> str:
    """Render a single alert as a structured text block for the system prompt."""
    ts          = alert.get("timestamp", "unknown")
    agent       = alert.get("agent_name", "unknown")
    rule_desc   = alert.get("rule_description", "—")
    rule_level  = alert.get("rule_level", 0)
    event_id    = alert.get("event_id", "—")
    channel     = alert.get("channel", "—")
    target_user = alert.get("target_user", "")
    subject_user= alert.get("subject_user", "")
    cmd_line    = alert.get("command_line", "")
    src_ip      = alert.get("src_ip", "")
    full_log    = alert.get("full_log", "")
    rule_groups = ", ".join(alert.get("rule_groups", [])) or "—"
    mitre_ids   = ", ".join(alert.get("mitre_ids", [])) or "—"
    mitre_tacs  = ", ".join(alert.get("mitre_tactics", [])) or "—"
    mitre_techs = ", ".join(alert.get("mitre_techniques", [])) or "—"

    lines = [f"--- {'ANOMALY' if is_anomaly else 'NORMAL'} #{index} ---"]
    lines.append(f"Timestamp    : {ts}")
    lines.append(f"Agent        : {agent}")
    lines.append(f"Rule         : [{rule_level}] {rule_desc}")
    lines.append(f"Event ID     : {event_id}  |  Channel: {channel}")
    lines.append(f"Rule Groups  : {rule_groups}")
    lines.append(f"MITRE IDs    : {mitre_ids}")
    lines.append(f"MITRE Tactics: {mitre_tacs}")
    lines.append(f"MITRE Techs  : {mitre_techs}")

    if target_user:
        lines.append(f"Target User  : {target_user}")
    if subject_user:
        lines.append(f"Subject User : {subject_user}")
    if src_ip:
        lines.append(f"Source IP    : {src_ip}")
    if cmd_line:
        lines.append(f"Command Line : {cmd_line[:300]}{'…' if len(cmd_line) > 300 else ''}")

    if is_anomaly:
        score      = alert.get("anomaly_score", 0.0)
        confidence = alert.get("confidence", "—")
        cluster_id = alert.get("cluster_id", None)
        mitre_stage= alert.get("mitre_stage", "—")
        shap_vals  = alert.get("shap_values", [])

        lines.append(f"Anomaly Score: {score:.4f} ({(score*100):.1f}%) | Confidence: {confidence}")
        if cluster_id is not None:
            lines.append(f"Cluster ID   : {cluster_id}")
        lines.append(f"MITRE Stage  : {mitre_stage or '—'}")

    if full_log:
        truncated = full_log[:400]
        lines.append(f"Raw Log      : {truncated}{'…' if len(full_log) > 400 else ''}")

    return "\n".join(lines)


def build_system_prompt(pipeline_data: dict | None) -> str:
    """
    Build the full system prompt fed to Gemini.
    Includes: persona, pipeline metadata, anomaly details, normal baselines,
    and SHAP-based correlation hints.
    """

    base_persona = (
        "You are ATHEA, an expert AI security analyst embedded in the Wazuh-ATHEA dashboard. "
        "Your role is to help SOC analysts understand security anomalies detected by an ensemble "
        "ML model.\n\n"
        "When answering:\n"
        "- Be specific and reference actual data from the alerts provided below.\n"
        "- Provide your own expert opinion on whether an anomaly is a genuine threat or a false positive.\n"
        "- If an anomaly appears to be a false positive (e.g., routine system updates, normal administrative tasks), explicitly state your reasoning.\n"
        "- Do NOT output or rely on SHAP values. Base your analysis on raw event data, rule descriptions, logs, and your own cybersecurity knowledge.\n"
        "- Correlate anomalies with normal logs to highlight what makes the flagged events "
        "stand out from the baseline.\n"
        "- Map findings to MITRE ATT&CK tactics when available.\n"
        "- If no pipeline data is available yet, say so and ask the user to wait for the "
        "first analysis cycle to complete.\n"
        "- Format responses in clear markdown with headers, bullet points, and code blocks "
        "where appropriate.\n"
    )

    if not pipeline_data:
        return (
            base_persona
            + "\n\n[CONTEXT]: No pipeline data is available yet. "
            "The first analysis cycle may still be running. "
            "Inform the user and advise them to wait a moment before retrying."
        )

    # ── Pipeline metadata ───────────────────────────────────────────────────
    total_alerts   = pipeline_data.get("total_alerts", 0)
    anomaly_count  = pipeline_data.get("anomaly_count", 0)
    normal_count   = pipeline_data.get("normal_count", 0)
    model_used     = pipeline_data.get("model_used", "unknown")
    contamination  = pipeline_data.get("contamination_used", 0.0)
    minutes_back   = pipeline_data.get("minutes_back", 60)
    pipeline_ver   = pipeline_data.get("pipeline_version", "—")
    conf_dist      = pipeline_data.get("confidence_distribution", {})
    top_mitre      = pipeline_data.get("top_mitre_stages", [])

    conf_str = ", ".join(f"{k}: {v}" for k, v in conf_dist.items()) or "none"
    mitre_str = ", ".join(f"{stage} ({cnt})" for stage, cnt in top_mitre) or "none"

    meta_block = (
        f"=== PIPELINE METADATA ===\n"
        f"Pipeline Version : {pipeline_ver}\n"
        f"Analysis Window  : Last {minutes_back} minutes\n"
        f"ML Model         : {model_used}\n"
        f"Contamination    : {contamination:.1%} (sensitivity setting)\n"
        f"Total Alerts     : {total_alerts}\n"
        f"Anomalies Flagged: {anomaly_count}\n"
        f"Normal Alerts    : {normal_count}\n"
        f"Confidence Dist  : {conf_str}\n"
        f"Top MITRE Stages : {mitre_str}\n"
    )

    # ── Anomalies block ─────────────────────────────────────────────────────
    all_alerts = pipeline_data.get("alerts", [])
    anomalies  = [a for a in all_alerts if a.get("anomaly") == 1]
    normals    = [a for a in all_alerts if a.get("anomaly") != 1]

    # Sort anomalies by score descending, take top 20
    anomalies_sorted = sorted(
        anomalies,
        key=lambda a: a.get("anomaly_score", 0.0),
        reverse=True
    )[:20]

    anomaly_block_lines = [f"\n=== ANOMALOUS ALERTS ({len(anomalies)} total, showing top {len(anomalies_sorted)}) ==="]
    for i, alert in enumerate(anomalies_sorted, 1):
        anomaly_block_lines.append(_format_alert_for_prompt(alert, is_anomaly=True, index=i))

    # ── Normal baselines block ───────────────────────────────────────────────
    # Take a representative sample of normals — spread across agents
    agent_seen: set = set()
    normals_sample: list[dict] = []
    for alert in normals:
        agent = alert.get("agent_name", "unknown")
        if agent not in agent_seen or len(normals_sample) < 5:
            normals_sample.append(alert)
            agent_seen.add(agent)
        if len(normals_sample) >= 10:
            break

    normal_block_lines = [f"\n=== NORMAL (BASELINE) ALERTS ({normal_count} total, showing {len(normals_sample)} representative samples) ==="]
    normal_block_lines.append(
        "(These events were NOT flagged as anomalous — they represent typical/expected activity "
        "and serve as the baseline for comparison.)"
    )
    for i, alert in enumerate(normals_sample, 1):
        normal_block_lines.append(_format_alert_for_prompt(alert, is_anomaly=False, index=i))

    # ── Correlation hint ─────────────────────────────────────────────────────
    correlation_hint = (
        "\n=== CORRELATION GUIDANCE ===\n"
        "When explaining anomalies, compare them against the normal baseline above. "
        "For example:\n"
        "  - If a normal alert has RuleLevel 3 and the anomaly has RuleLevel 14, explain the difference.\n"
        "  - Check the raw log and command line for routine administrative tasks (like apt/dpkg installations) that might be false positives.\n"
        "  - If LogEntropy is high in the anomaly but low in normals, suggest possible obfuscation.\n"
        "  - Reference cluster IDs to group related anomalies into attack campaigns.\n"
        "  - Use MITRE stages to narrate the attack kill chain if multiple stages appear.\n"
    )

    full_prompt = (
        base_persona
        + "\n\n"
        + meta_block
        + "\n".join(anomaly_block_lines)
        + "\n"
        + "\n".join(normal_block_lines)
        + "\n"
        + correlation_hint
        + "\n\n"
        + "Answer the analyst's question below using the data provided above. "
        "Be precise, cite specific alert details, and provide your own expert opinion on the threat level.\n"
    )

    return full_prompt


def build_report_prompt(pipeline_data: dict | None) -> str:
    """Build the prompt for generating the Security Anomaly Intelligence Report."""
    if not pipeline_data:
        return "You are an AI security analyst. The user requested a report, but no data is available yet."

    import datetime
    import random
    import string
    now = datetime.datetime.now()
    date_str = now.strftime("%Y%m%d")
    time_str = now.strftime("%H%M")
    random_str = ''.join(random.choices(string.ascii_uppercase + string.digits, k=4))
    report_id = f"{date_str}{random_str}{time_str}"

    minutes = pipeline_data.get("minutes_back", 60)
    start_time = now - datetime.timedelta(minutes=minutes)
    time_window = f"{start_time.strftime('%H:%M')} - {now.strftime('%H:%M')}"
    pipeline_ver = pipeline_data.get("pipeline_version", "unknown")
    
    # Use the same formatting logic as the system prompt to get context

    # Use the same formatting logic as the system prompt to get context
    base_context = build_system_prompt(pipeline_data)

    report_instructions = f"""
Generate a formal "Security Anomaly Intelligence Report" exactly matching the structure below. Focus entirely on providing in-depth, professional explanations of the security context, attack narratives, and practical implications. Do NOT mention machine learning models, algorithms, anomaly scores, SHAP values, or training metrics. Explain the situation as an elite human security analyst would. Do NOT output markdown code blocks around the entire response.

Structure:
Security Anomaly Intelligence Report
Reporting Window: {time_window}
Report ID: {report_id}

1. Executive Summary
[Provide a comprehensive, high-level professional summary of the security situation. Use descriptive bullet points to thoroughly explain the real-world impact, the nature of the detected threats, and their broader significance to the organization. Do not use model metrics.]

2. Threat Narrative & MITRE ATT&CK
[Provide an in-depth explanation of the potential attack narrative. Use bullet points to map observed behaviors to MITRE ATT&CK stages (Tactics and Techniques). Professionally explain *how* these techniques work together and *what* the adversary's ultimate objective likely is.]

3. Technical & Behavioral Deep-Dive
[Pick the most critical anomaly. Use descriptive bullet points to provide extensive technical details. Explicitly list out specific Event IDs, process names, command lines, IP addresses, user accounts, and file paths involved. For every technical detail provided, include a professional explanation of its function and why it is suspicious in this context.]

4. Plain-English Explanation
[Provide a highly detailed, professional explanation of the deep-dive anomaly using bullet points: Explain exactly what this activity means in plain English. Why did the system generate this alert? Thoroughly explain under what circumstances this behavior represents a critical threat, and contrast it with scenarios where it might be a routine IT administrative operation.]

5. Analyst Observations & False Positive Analysis
[Provide your expert, detailed assessment using professional pointers. Thoroughly explain whether these are likely true positive threats or false positives. Provide deep reasoning based on the system logs, user context, and typical enterprise environments, avoiding any reference to model confidence.]

6. Recommended Actions
[Provide highly detailed, actionable, and concrete steps for investigation, mitigation, and remediation formatted as a clear bulleted list. Explain *why* each action is recommended.]

FORMATTING RULES:
- Use bullet points extensively across all sections. Ensure each bullet point contains a thorough, multi-sentence professional explanation rather than just a brief note.
- ONLY use bold (**text**) for the main numbered headings (e.g., **1. Executive Summary**) and key technical terms within bullet points (e.g., **Event ID 4688**, **powershell.exe**).
- DO NOT use bold for entire sentences or paragraphs.
- DO NOT use italics for sentences, paragraphs, or general text analysis.
- ONLY use italics (*text*) strictly for short terminology, pointer names, keys, or feature names.
"""
    return base_context + "\n\n" + report_instructions



def _convert_history_to_gemini(history: list[ChatMessage]) -> list[dict[str, Any]]:
    """Convert our ChatMessage list to the format Gemini expects."""
    gemini_history = []
    for msg in history:
        # Gemini uses "model" for assistant role
        role = "model" if msg.role in ("model", "assistant") else "user"
        gemini_history.append({
            "role": role,
            "parts": [{"text": msg.content}],
        })
    return gemini_history


# ── FastAPI Router ───────────────────────────────────────────────────────────
router = APIRouter(tags=["Chatbot"])


@router.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """
    Send a message to the ATHEA Gemini chatbot.
    The chatbot is automatically provided the latest pipeline results as context.
    """
    # _state is injected via dependency — imported at call time from main.py
    from main import _state as app_state

    pipeline_data: dict | None = app_state.get("cached_results")

    if not _gemini_ready:
        raise HTTPException(
            status_code=503,
            detail="Gemini API key not configured. Please set GEMINI_API_KEY in your .env file.",
        )

    system_prompt = build_system_prompt(pipeline_data)
    context_loaded = pipeline_data is not None

    # Convert conversation history to Gemini format
    gemini_history = _convert_history_to_gemini(request.history)

    try:
        model = genai.GenerativeModel(
            model_name=GEMINI_MODEL,
            system_instruction=system_prompt,
        )

        chat_session = model.start_chat(history=gemini_history)
        response = chat_session.send_message(request.message)

        reply_text = response.text

    except Exception as exc:
        logger.error(f"Gemini API error: {exc}", exc_info=True)
        raise HTTPException(
            status_code=502,
            detail=f"Gemini API error: {str(exc)}",
        )

    return ChatResponse(
        reply=reply_text,
        model=GEMINI_MODEL,
        context_loaded=context_loaded,
    )


@router.post("/api/report", response_model=ReportResponse)
async def generate_report():
    """
    Generate a formal Security Anomaly Intelligence Report using Gemini.
    """
    from main import _state as app_state

    pipeline_data: dict | None = app_state.get("cached_results")

    if not _gemini_ready:
        raise HTTPException(
            status_code=503,
            detail="Gemini API key not configured. Please set GEMINI_API_KEY in your .env file.",
        )

    prompt = build_report_prompt(pipeline_data)

    try:
        model = genai.GenerativeModel(model_name=GEMINI_MODEL)
        response = model.generate_content(prompt)
        reply_text = response.text
    except Exception as exc:
        logger.error(f"Gemini API error during report generation: {exc}", exc_info=True)
        raise HTTPException(
            status_code=502,
            detail=f"Gemini API error: {str(exc)}",
        )

    return ReportResponse(report=reply_text)


def generate_docx_from_markdown(md_text: str) -> io.BytesIO:
    """Convert basic markdown text to a professionally styled DOCX file."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise RuntimeError("python-docx is not installed.")

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = md_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('Security Anomaly Intelligence Report'):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
            continue
            
        # Headings
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 51, 102)
            continue
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0, 51, 102)
            continue
        elif line.startswith('# ') or (len(line) > 2 and line[0].isdigit() and line[1:3] == '. '):
            # Also treat "1. Executive Summary" as a header
            p = doc.add_paragraph()
            if line.startswith('# '):
                run = p.add_run(line[2:])
            else:
                run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 51, 102)
            continue
            
        # Lists
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
        else:
            p = doc.add_paragraph()
            text = line
            
        # Parse bold (**text**) and italic (*text*)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                subparts = re.split(r'(\*[^*\n]+\*)', part)
                for subpart in subparts:
                    if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                        run = p.add_run(subpart[1:-1])
                        run.italic = True
                    else:
                        p.add_run(subpart)

    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f


@router.post("/api/report/docx")
async def download_docx(req: DocxRequest):
    """Generate a downloadable DOCX file from report text."""
    try:
        file_stream = generate_docx_from_markdown(req.report_text)
        from fastapi.responses import Response
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Intelligence_Report.docx"}
        )
    except Exception as exc:
        logger.error(f"Error generating DOCX: {exc}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(exc))
